from io import BytesIO

from docx import Document


def generate_cv_docx(cv):
    doc = Document()

    doc.add_heading(cv.user.username, level=1)
    doc.add_paragraph(cv.work_email)
    doc.add_paragraph(cv.phone_number)

    if cv.github:
        doc.add_paragraph(f"GitHub: {cv.github}")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(cv.summary)

    doc.add_heading("Technical skills", level=2)
    skills = ", ".join(skill.get_code_display() for skill in cv.skills.all())
    doc.add_paragraph(skills)

    doc.add_heading("Work experience", level=2)
    doc.add_paragraph(cv.work_experience)

    doc.add_heading("Education", level=2)
    doc.add_paragraph(cv.education)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
